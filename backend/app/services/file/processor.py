# app/services/file/processor.py
from pathlib import Path
from typing import Optional, Tuple
import magic
import pypdf
import pytesseract
from PIL import Image
from docx import Document
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.file import File
from app.schemas.file_processing import ProcessedFile
from app.services.rag.processor import RAGProcessor

import logging

logger = logging.getLogger(__name__)

class FileProcessor:
    """Handles file processing and text extraction"""

    def __init__(self, rag_processor: RAGProcessor):
        self.rag_processor = rag_processor
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)

    async def process_and_vectorize(
            self,
            processed_file: ProcessedFile,
            file_id: str,
            db: Session
    ) -> None:
        """Process file contents and generate vectors (runs in background)"""
        try:
            # Extract text and metadata
            metadata = {
                "title": processed_file.name,
                "mime_type": processed_file.mime_type,
                "size": processed_file.size,
                "pages": processed_file.page_count
            }
            logger.info(f"Processing file {file_id}")

            # Process through RAG pipeline
            status = await self.rag_processor.process_document(
                text=processed_file.content,
                metadata=metadata,
                file_info=processed_file
            )
            logger.info(f"RAG processing finished for file {file_id} with status {status.status}")

            if status.status == "completed":
                # Update database status
                file = db.query(File).filter(File.id == file_id).first()
                if file:
                    file.vectorized = True
                    db.commit()
                    logger.info(f"File {file_id} vectorized successfully")

        except Exception as e:
            # Log error and update database status
            print(f"Error processing file {file_id}: {str(e)}")
            file = db.query(File).filter(File.id == file_id).first()
            if file:
                file.vectorized = False
                db.commit()

    def _get_mime_type(self, file_content: bytes) -> str:
        """Detect file MIME type"""
        mime = magic.Magic(mime=True)
        return mime.from_buffer(file_content[:2048])

    def _extract_text_from_pdf(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from PDF file"""
        with open(file_path, 'rb') as file:
            pdf = pypdf.PdfReader(file)
            text_parts = []
            for page in pdf.pages:
                text_parts.append(page.extract_text())
            return '\n'.join(text_parts), len(pdf.pages)

    def _extract_text_from_docx(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        return '\n'.join(text_parts), len(doc.paragraphs)

    def _extract_text_from_image(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from image using OCR"""
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text, 1

    async def process_file(self, file: UploadFile) -> ProcessedFile:
        """Initial file processing and text extraction"""
        if file.size > settings.MAX_FILE_SIZE:
            raise HTTPException(400, "File too large")

        # Save file temporarily
        file_path = self.upload_dir / file.filename
        logger.info(f"Saving file to {file_path}")
        try:
            content = await file.read()
            mime_type = self._get_mime_type(content)

            with open(file_path, 'wb') as f:
                f.write(content)

            # Extract text based on file type
            text, page_count = await self._extract_text(file_path, mime_type)
            logger.info(f"Text extracted from {file_path}")

            return ProcessedFile(
                id=str(file_path.stem),
                name=file.filename,
                size=file.size,
                mime_type=mime_type,
                page_count=page_count,
                content=text,
                metadata={}
            )

        finally:
            logger.info(f"Cleaning up")
            if file_path.exists():
                file_path.unlink()

    async def _extract_text(
            self,
            file_path: Path,
            mime_type: str
    ) -> Tuple[str, int]:
        """Extract text from different file types"""
        try:
            logger.info(f"Extracting text")
            if mime_type == 'application/pdf':
                logger.info(f"Extracting text from {file_path}")
                return self._extract_text_from_pdf(file_path)
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                logger.info(f"Extracting text from {file_path}")
                return self._extract_text_from_docx(file_path)
            elif mime_type.startswith('image/'):
                logger.info(f"Extracting text from {file_path}")
                return self._extract_text_from_image(file_path)
            else:
                raise HTTPException(400, f"Unsupported file type: {mime_type}")
        except Exception as e:
            raise HTTPException(500, f"Text extraction failed: {str(e)}")